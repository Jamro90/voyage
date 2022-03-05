from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ROW_HEIGHT_RULE
from tkinter import *
from tkinter.ttk import *

def sort_list(pchor_list, from_date, to_date, check_var):
    pchor_list_mod = []
    i = 0
    # deleting labels row
    pchor_list.pop(0)
    for pchor in pchor_list:
        # changing string in to array
        line = pchor.split(",")
        # removing indexing
        line.append(from_date[i].get())
        line.append(to_date[i].get())
        line.append(check_var[i].get())
        line.pop(0)
        pchor_list_mod.append(line)
        i += 1

    pchor_list_moded = []
    for pchor in pchor_list_mod:
        if pchor[4] == 1:
            pchor_list_moded.append(pchor)

    # sorting levels
    pchor_list_moded.sort()
    # making levels arrays
    sergant_list = []
    platoon_sergant_list = []
    corporal_1_list = []
    corporal_list = []
    private_1_list = []
    private_list = []

    # sorting levels
    for pchor in pchor_list_moded:
        if pchor[0] == "sierż. pchor.":
            sergant_list.append(pchor)
        
        elif pchor[0] == "plut. pchor.":
            platoon_sergant_list.append(pchor)
        
        elif pchor[0] == "st. kpr. pchor.":
            corporal_1_list.append(pchor)
        
        elif pchor[0] == "kpr. pchor.":
            corporal_list.append(pchor)
            
        elif pchor[0] == "st. szer. pchor.":
            private_1_list.append(pchor)
            
        elif pchor[0] == "szer. pchor.":
            private_list.append(pchor)
        
        else:
            print("nie działa")

    # making correct indexing
    i = 1
    for sergant in sergant_list:
        sergant.insert(0, i)
        i += 1
    for platoon_sergant in platoon_sergant_list:
        platoon_sergant.insert(0, i)
        i += 1
    for corporal_1 in corporal_1_list:
        corporal_1.insert(0, i)
        i += 1
    for corporal in corporal_list:
        corporal.insert(0, i)
        i += 1
    for private_1 in private_1_list:
        private_1.insert(0, i)
        i += 1
    for private in private_list:
        private.insert(0, i)
        i += 1
    
    # finale sorting algorithm
    num_sorted_list = []
    for sergant in sergant_list:
        num_sorted_list.append(sergant)
    for platoon_sergant in platoon_sergant_list:
        num_sorted_list.append(platoon_sergant)
    for corporal_1 in corporal_1_list:
        num_sorted_list.append(corporal_1)
    for corporal in corporal_list:
        num_sorted_list.append(corporal)
    for private_1 in private_1_list:
        num_sorted_list.append(private_1)
    for private in private_list:
        num_sorted_list.append(private)
    
    return num_sorted_list

def doc_make(sorted_data, window):
    file_name = filedialog.asksaveasfile(parent = window, title = "Zapisz", filetypes = (("MS WORD 97", "*.doc"), ("MS WORD", "*.docx"), ("Writer", "*.odt")))
    
    font_name = "Times New Roman"
    font_size = 12
    
    name = []
    surname = []
    i = 0
    for d in sorted_data:
        split = sorted_data[i][2].split(" ")
        name.append(split[1])
        surname.append(split[0])
        i += 1
    # making document
    document = Document()
    # title text
    title = document.add_paragraph()
    title_text = title.add_run("7.    PRZEPUSTKI JEDNORAZOWE:")
    title_format = title_text.font
    title_format.name = font_name
    title_format.size = font_size
    title_format.bold = True
    # sub title text
    sub_title = document.add_paragraph()
    sub_title_text = sub_title.add_run("a)    z 2 KMP")
    sub_title_format = sub_title_text.font
    sub_title_format.name = font_name
    sub_title_format.size = font_size
    sub_title_format.bold = True
    sub_title.paragraph_format.line_spacing = 1.15
    # loop for writing records
    i = 0
    table = document.add_table(rows = len(sorted_data), cols = 6)
    table.columns[0].width = Inches(0.5)
    table.columns[1].width = Inches(1.2)
    table.columns[2].width = Inches(1.3)
    table.columns[3].width = Inches(1.3)
    table.columns[4].width = Inches(0.6)
    table.columns[5].width = Inches(2)
    for data in sorted_data:
        table.rows[i].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[i].height = Inches(0.25)
        index_cell = table.cell(i, 0)
        level_cell = table.cell(i, 1)
        name_cell = table.cell(i, 2)
        surname_cell = table.cell(i, 3)
        null_cell = table.cell(i, 4)
        date_cell = table.cell(i, 5)

        index_cell.text = f"{data[0]})"
        level_cell.text = f"{data[1]}"
        name_cell.text = f"{name[i]}"
        surname_cell.text = f"{surname[i].upper()}"
        null_cell.text = "w dn."
        date_cell.text = f"{data[3][:2]}-{data[4]}"
        i += 1
    # save as reqested file (including data extension)
    document.save(file_name.name)
